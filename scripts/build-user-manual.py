#!/usr/bin/env python3
"""Build the HS Voice v1.2.0 Japanese end-user manual."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Design system: compact_reference_guide, with a named Japanese font override.
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

# Named preset override: use one Unicode family for every script.
# LibreOffice did not expose Japanese glyphs from Hiragino during headless export.
LATIN_FONT = "Arial Unicode MS"
JAPANESE_FONT = "Arial Unicode MS"
INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "11AFA4"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_TEAL = "E9F8F6"
CAUTION_FILL = "FFF8E8"
CAUTION = "7A5A00"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in CELL_MARGINS_DXA.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for element in list(tbl_pr.findall(qn(tag))):
            tbl_pr.remove(element)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row) -> None:
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(tbl_header)


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = LATIN_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{attribute}"), None)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), JAPANESE_FONT)
    r_fonts.set(qn("w:cs"), JAPANESE_FONT)
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "ja-JP")
    language.set(qn("w:eastAsia"), "ja-JP")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, *, size, color, before, after, line_spacing, bold=False) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{attribute}"), None)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), JAPANESE_FONT)
    r_fonts.set(qn("w:cs"), JAPANESE_FONT)
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "ja-JP")
    language.set(qn("w:eastAsia"), "ja-JP")
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    paragraph.keep_with_next = style.name.startswith("Heading")


def set_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_document_defaults(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.insert(0, r_pr_default)
    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.attrib.clear()
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), JAPANESE_FONT)
    r_fonts.set(qn("w:cs"), JAPANESE_FONT)
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "ja-JP")
    language.set(qn("w:eastAsia"), "ja-JP")
    language.set(qn("w:bidi"), "ja-JP")


def paragraph_bottom_border(paragraph, *, color="D0D5DD", size="6", space="5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def configure_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("HS Voice  |  利用マニュアル")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    paragraph_bottom_border(paragraph, color="D0D5DD", size="4", space="4")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(2)
    label = paragraph.add_run("HS Voice 1.2.0   |   ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_field(paragraph, "PAGE")
    separator = paragraph.add_run(" / ")
    set_run_font(separator, size=8.5, color=MUTED)
    add_page_field(paragraph, "NUMPAGES")


def add_numbering_definition(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    if not ordered:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(fonts)
        level.append(r_pr)

    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_paragraph(doc, text="", *, bold_prefix=None, color=None, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True, color=color or INK)
        remainder = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(remainder, color=color or INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=color or INK)
    return paragraph


def add_bullets(doc, items: list[str]) -> None:
    num_id = add_numbering_definition(doc, ordered=False)
    for item in items:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        set_run_font(run, color=INK)


def add_steps(doc, items: list[tuple[str, str]]) -> None:
    num_id = add_numbering_definition(doc, ordered=True)
    for title, detail in items:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.25
        title_run = paragraph.add_run(title)
        set_run_font(title_run, bold=True, color=INK)
        detail_run = paragraph.add_run(f"  {detail}")
        set_run_font(detail_run, color=INK)


def add_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_callout(doc, label: str, text: str, *, kind="info") -> None:
    fills = {"info": PALE_TEAL, "note": LIGHT_GRAY, "caution": CAUTION_FILL}
    label_colors = {"info": TEAL, "note": DARK_BLUE, "caution": CAUTION}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fills[kind])
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge, size in (("left", "18"), ("top", "3"), ("bottom", "3"), ("right", "3")):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), label_colors[kind])
        borders.append(border)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, bold=True, color=label_colors[kind])
    body_run = paragraph.add_run(text)
    set_run_font(body_run, color=INK)


def format_table_text(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=10,
                        bold=row_index == 0,
                        color=WHITE if row_index == 0 else INK,
                    )


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        set_cell_shading(table.cell(0, index), BLUE)
    mark_header_row(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_geometry(table, widths_dxa)
    format_table_text(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_break(doc) -> None:
    doc.add_page_break()


def add_cover(doc: Document, icon_path: Path) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(64)

    icon_paragraph = doc.add_paragraph()
    icon_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = icon_paragraph.add_run()
    inline = run.add_picture(str(icon_path), width=Inches(1.25))
    inline._inline.docPr.set("descr", "HS Voiceの青緑色の波形アプリアイコン")
    icon_paragraph.paragraph_format.space_after = Pt(20)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("社内向けスマート音声入力")
    set_run_font(run, size=11, bold=True, color=TEAL)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(9)
    run = title.add_run("HS Voice")
    set_run_font(run, size=31, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("利用マニュアル")
    set_run_font(run, size=20, bold=True, color=BLUE)

    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.space_after = Pt(76)
    statement.paragraph_format.line_spacing = 1.25
    run = statement.add_run("話すだけで、どこにでも入力。\nMacBookでの初期設定から日常操作まで")
    set_run_font(run, size=12.5, color=MUTED)

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(4)
    run = metadata.add_run("対象バージョン  1.2.0  |  macOS 14以降")
    set_run_font(run, size=10.5, bold=True, color=INK)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run("2026年8月版")
    set_run_font(run, size=9.5, color=MUTED)


def build_manual(output_path: Path, icon_path: Path) -> None:
    doc = Document()
    configure_document_defaults(doc)
    for section in doc.sections:
        set_section_geometry(section)
        configure_header_footer(section)

    styles = doc.styles
    configure_style(
        styles["Normal"], size=11, color=INK, before=0, after=6, line_spacing=1.25
    )
    configure_style(
        styles["Heading 1"], size=16, color=BLUE, before=18, after=10,
        line_spacing=1.0, bold=True
    )
    configure_style(
        styles["Heading 2"], size=13, color=BLUE, before=14, after=7,
        line_spacing=1.0, bold=True
    )
    configure_style(
        styles["Heading 3"], size=12, color=DARK_BLUE, before=10, after=5,
        line_spacing=1.0, bold=True
    )

    core = doc.core_properties
    core.title = "HS Voice 1.2.0 利用マニュアル"
    core.subject = "社内向けmacOS音声入力アプリの利用方法"
    core.comments = "HS Voice 1.2.0 社内向け利用マニュアル"
    core.keywords = "HS Voice, macOS, 音声入力, 利用マニュアル"
    core.author = "HS Voice"
    core.last_modified_by = "HS Voice"
    core.created = datetime(2026, 8, 24)
    core.modified = datetime(2026, 8, 24)

    add_cover(doc, icon_path)
    add_page_break(doc)

    add_heading(doc, "1. はじめに", 1)
    add_paragraph(
        doc,
        "HS Voiceは、Macのメニューバーに常駐し、話した内容をテキストへ変換する社内向け音声入力アプリです。メール、チャット、文書作成など、カーソルを置ける多くのアプリで利用できます。",
    )
    add_callout(
        doc,
        "最短の使い方",
        "入力したい欄へカーソルを置き、fn（地球儀）キーを押しながら話し、話し終わったらキーを離します。",
    )

    add_heading(doc, "動作環境", 2)
    add_bullets(
        doc,
        [
            "macOS 14 Sonoma以降",
            "Apple Siliconまたは64-bit Intel搭載Mac",
            "利用可能なマイクとインターネット接続（認識言語や端末の状態により必要）",
        ],
    )

    add_heading(doc, "必要な権限", 2)
    add_table(
        doc,
        ["権限", "必要性", "用途"],
        [
            ["マイク", "必須", "声を取り込む"],
            ["音声認識", "必須", "話した内容をテキストへ変換する"],
            ["アクセシビリティ", "自動入力時のみ", "認識結果を他のアプリへ貼り付ける"],
        ],
        [2050, 1970, 5340],
    )
    add_callout(
        doc,
        "権限を限定したい場合",
        "アクセシビリティを許可せず、「クリップボードへコピーのみ」で利用できます。",
        kind="note",
    )

    add_heading(doc, "このマニュアルの構成", 2)
    add_bullets(
        doc,
        [
            "初回セットアップと基本操作",
            "メニューバーと直前テキストの再利用",
            "設定、プライバシー、履歴",
            "困ったときの確認方法と社内ITへの問い合わせ",
        ],
    )
    add_page_break(doc)

    add_heading(doc, "2. インストールと初回セットアップ", 1)
    add_heading(doc, "インストール", 2)
    add_steps(
        doc,
        [
            ("配布されたPKGを開く。", "会社の案内に従い、HSVoice-Installer-1.2.0.pkgをダブルクリックします。"),
            ("インストーラを進める。", "画面の案内に従い、HS VoiceをApplicationsフォルダへインストールします。"),
            ("HS Voiceを起動する。", "Finderの「アプリケーション」からHS Voiceを開きます。起動後はメニューバーに波形アイコンが表示されます。"),
        ],
    )
    add_callout(
        doc,
        "セキュリティ注意",
        "「開発元を確認できません」など会社の配布手順と異なる警告が出た場合は、自己判断で回避せず社内ITへ連絡してください。",
        kind="caution",
    )

    add_heading(doc, "初回セットアップ", 2)
    add_steps(
        doc,
        [
            ("セットアップを開始する。", "「セットアップを開始」を1回押します。HS Voice独自の設定入力はありません。"),
            ("マイクと音声認識を許可する。", "続けて表示されるmacOSの確認で、それぞれ「許可」を選びます。"),
            ("自動入力を許可する。", "開いた「プライバシーとセキュリティ」>「アクセシビリティ」でHS Voiceを有効にします。"),
            ("そのまま利用を開始する。", "HS Voiceへ戻るとガイドは自動で閉じます。許可しない場合は「コピーだけで使う」を選べます。"),
        ],
    )
    add_paragraph(
        doc,
        "システム設定からHS Voiceへ戻ると、許可状態は自動更新されます。反映されない場合は、メニューバーの「権限」または設定の「権限」タブから状態を更新してください。",
    )
    add_page_break(doc)

    add_heading(doc, "3. 基本の音声入力", 1)
    add_heading(doc, "押している間だけ録音する（初期設定）", 2)
    add_table(
        doc,
        ["手順", "操作", "内容"],
        [
            ["1", "入力先を選ぶ", "メールやチャットなどのテキスト欄をクリックし、カーソルを表示します。"],
            ["2", "fnキーを押し続ける", "画面下部に録音オーバーレイが表示されます。"],
            ["3", "自然に話す", "音声レベル、経過時間、認識途中のテキストを確認できます。"],
            ["4", "キーを離す", "認識結果が整形され、自動入力またはクリップボードへのコピーが行われます。"],
        ],
        [900, 2900, 5560],
    )
    add_callout(
        doc,
        "うまく認識させるコツ",
        "マイクから一定の距離で、短い文ごとにはっきり話します。1回の入力は55秒で自動終了するため、長文は段落ごとに分けてください。",
    )

    add_heading(doc, "トグル式で録音する", 2)
    add_paragraph(
        doc,
        "設定の「一般」>「起動方法」で「押すたびに開始・停止」を選ぶと、ショートカットを1回押して録音開始、もう1回押して終了できます。録音を止め忘れても55秒で安全に終了します。",
    )

    add_heading(doc, "音声レイアウトコマンド", 2)
    add_table(
        doc,
        ["話す言葉", "入力結果", "利用例"],
        [
            ["改行", "1行改行", "箇条書きに近い短い行を続ける"],
            ["新しい段落", "空行を含む段落区切り", "話題や章を切り替える"],
            ["new line", "1行改行", "英語入力で行を分ける"],
            ["new paragraph", "空行を含む段落区切り", "英語入力で段落を分ける"],
        ],
        [2350, 2210, 4800],
    )
    add_paragraph(
        doc,
        "音声コマンドを文字どおり入力したい場合は、設定の「言語と辞書」で音声コマンド整形をオフにします。",
    )
    add_page_break(doc)

    add_heading(doc, "4. メニューバーを使いこなす", 1)
    add_paragraph(
        doc,
        "メニューバーの波形アイコンをクリックすると、録音、言語変更、入力方法、直前テキストの操作へすぐアクセスできます。",
    )
    add_table(
        doc,
        ["画面の場所", "できること"],
        [
            ["上部の状態表示", "待機中、聞いています、仕上げています、入力しました、確認が必要です、の現在状態を確認する"],
            ["大きな録音カード", "クリックして録音を開始・終了する。ショートカットと同じ操作"],
            ["話す言語", "設定画面を開かずに、日本語や英語などへ切り替える"],
            ["入力方法", "自動入力とコピーのみを切り替える"],
            ["直前の入力", "もう一度入力、コピー、条件を満たす場合は取り消す"],
            ["下部メニュー", "履歴、設定、権限、終了へ移動する"],
        ],
        [2700, 6660],
    )

    add_heading(doc, "直前のテキストを再利用する", 2)
    add_bullets(
        doc,
        [
            "もう一度入力: 現在の入力方法で直前のテキストを再度処理します。",
            "コピー: 直前のテキストをクリップボードへコピーします。",
            "取り消す: 自動入力成功後8秒以内かつ同じ入力先アプリが前面にある場合だけ、その入力を取り消します。",
        ],
    )
    add_callout(
        doc,
        "取り消しの安全制限",
        "入力先が変わった場合や8秒を過ぎた場合、取り消しは表示されないか実行されません。必要に応じて入力先アプリ側の取り消し操作を使用してください。",
        kind="note",
    )

    add_heading(doc, "入力方法の違い", 2)
    add_table(
        doc,
        ["入力方法", "動作", "向いている場面"],
        [
            ["カーソル位置へ自動入力", "録音前に使っていたアプリを前面に戻し、認識結果を貼り付ける", "日常的なメール、チャット、文書作成"],
            ["クリップボードへコピーのみ", "自動貼り付けを行わず、結果をクリップボードへ残す", "制限のある業務アプリ、アクセシビリティを許可しない運用"],
        ],
        [2500, 3500, 3360],
    )
    add_page_break(doc)

    add_heading(doc, "5. 設定を変更する", 1)
    add_heading(doc, "一般", 2)
    add_table(
        doc,
        ["項目", "内容"],
        [
            ["起動方法", "押している間だけ録音、または押すたびに開始・停止"],
            ["入力方法", "自動入力、またはクリップボードへコピーのみ"],
            ["グローバルショートカット", "fn（初期設定）または4種類のSpace系キーから選択"],
            ["ログイン時に起動", "Macへのログイン後にHS Voiceを自動起動"],
            ["診断情報をコピー", "社内ITへ渡すための設定・権限情報をコピー"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "選べるショートカット", 2)
    add_table(
        doc,
        ["キー", "備考"],
        [
            ["fn（地球儀）", "初期設定。アクセシビリティ権限を使用"],
            ["Option + Space", "fnキーがない外付けキーボード向けの候補"],
            ["Control + Space", "macOSや入力ソースの設定と競合する場合は別のキーを選択"],
            ["Command + Shift + Space", "3キーの組み合わせ"],
            ["Control + Option + Space", "他アプリとの競合を避けたい場合の候補"],
        ],
        [3500, 5860],
    )

    add_heading(doc, "言語と辞書", 2)
    add_bullets(
        doc,
        [
            "話す言語: 日本語、英語、中国語、韓国語、フランス語、ドイツ語、スペイン語、イタリア語、ポルトガル語から選択できます。",
            "オンデバイス認識: 対応している言語とMacでは、オンデバイス処理を優先します。",
            "音声コマンド整形: 「改行」「新しい段落」などをレイアウトへ変換します。",
            "カスタム辞書: 人名、製品名、専門用語を1行に1つ、最大100語まで登録できます。",
        ],
    )
    add_page_break(doc)

    add_heading(doc, "6. 履歴とプライバシー", 1)
    add_heading(doc, "履歴", 2)
    add_paragraph(
        doc,
        "音声入力の履歴は初期設定でオフです。設定の「プライバシー」で有効にすると、認識テキストをこのMac内へ最大100件保存します。音声ファイルは保存しません。",
    )
    add_bullets(
        doc,
        [
            "メニューバーの「履歴」で保存済みテキストを確認できます。",
            "設定から履歴をすべて削除できます。",
            "機密情報を扱う場合は、会社のルールに従い履歴をオフのまま利用してください。",
        ],
    )

    add_heading(doc, "データの扱い", 2)
    add_table(
        doc,
        ["データ", "HS Voiceでの扱い"],
        [
            ["音声", "ファイルとして保存しない"],
            ["認識テキスト", "履歴を有効にした場合だけ最大100件をローカル保存"],
            ["カスタム辞書と設定", "このMac内のユーザー設定として保存"],
            ["音声認識処理", "環境によりAppleの音声認識サービスへ接続する場合がある"],
            ["診断情報", "明示的にコピーしたときだけクリップボードへ出力"],
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "診断情報に含まれないもの",
        "音声、入力本文、カスタム辞書、ユーザー名、端末名は含まれません。",
    )

    add_heading(doc, "業務利用時の注意", 2)
    add_bullets(
        doc,
        [
            "音声入力を禁止されている情報や場所では使用しないでください。",
            "入力後は、宛先、数字、人名、製品名、機密区分を送信前に確認してください。",
            "完全なオフライン動作を保証する機能ではありません。",
        ],
    )
    add_page_break(doc)

    add_heading(doc, "7. 困ったとき", 1)
    add_table(
        doc,
        ["症状", "確認すること"],
        [
            ["録音が始まらない", "マイクと音声認識が許可済みか、設定の「権限」で確認する"],
            ["文字が自動入力されない", "アクセシビリティを許可するか、「コピーのみ」で結果をCommand + Vして貼り付ける"],
            ["別のアプリへ入力されそう", "録音前に正しい入力欄をクリックする。入力先を特定できない場合はコピーのみへ切り替える"],
            ["ショートカットが動かない", "fnの場合はアクセシビリティを確認する。外付けキーボードではSpace系キーへ変更する"],
            ["認識精度が低い", "話す言語を確認し、カスタム辞書へ固有名詞を登録する。短い文に分ける"],
            ["ネットワークエラーになる", "社内ネットワークを確認する。オンデバイス認識に対応しない場合は接続が必要"],
            ["取り消しが表示されない", "取り消しは自動入力成功後8秒以内、同じ入力先が前面にある場合だけ利用可能"],
            ["権限変更が反映されない", "HS Voiceへ戻るか「状態を更新」を押す。それでも直らない場合はアプリを再起動"],
        ],
        [2780, 6580],
    )

    add_heading(doc, "社内ITへ問い合わせる", 2)
    add_steps(
        doc,
        [
            ("問題を再現する。", "発生したアプリ、操作、表示されたメッセージを確認します。機密の入力本文は送らないでください。"),
            ("診断情報をコピーする。", "メニューバーの「設定」>「一般」>「診断情報をコピー」を押します。"),
            ("社内ITへ送る。", "診断情報と、問題が起きた時刻・入力先アプリ名を社内の指定窓口へ送ります。"),
        ],
    )
    add_callout(
        doc,
        "診断情報の確認",
        "送信前にクリップボード内容を確認できます。診断情報には認識テキストや音声は含まれません。",
        kind="note",
    )
    add_page_break(doc)

    add_heading(doc, "8. クイックリファレンス", 1)
    add_callout(
        doc,
        "毎日の基本操作",
        "カーソルを置く  >  fnキーを押しながら話す  >  キーを離す  >  結果を確認する",
    )

    add_heading(doc, "操作早見表", 2)
    add_table(
        doc,
        ["したいこと", "操作"],
        [
            ["音声入力を開始・終了", "fnキーを押したまま話し、離す"],
            ["言語をすぐ変更", "メニューバー > 話す言語"],
            ["コピー運用へ変更", "メニューバー > 入力方法 > コピーのみ"],
            ["直前の文章を再入力", "メニューバー > 直前の入力 > もう一度入力"],
            ["直前の文章をコピー", "メニューバー > 直前の入力 > コピー"],
            ["入力直後に取り消す", "8秒以内にメニューバー > 取り消す"],
            ["履歴を見る", "メニューバー > 履歴（履歴保存を有効にしている場合）"],
            ["設定・権限を確認", "メニューバー > 設定、または権限"],
        ],
        [3300, 6060],
    )

    add_heading(doc, "入力前チェック", 2)
    add_bullets(
        doc,
        [
            "入力先の欄にカーソルがある",
            "話す言語が合っている",
            "自動入力またはコピーのみの選択が用途に合っている",
            "機密情報の取り扱いが会社のルールに合っている",
        ],
    )

    add_heading(doc, "入力後チェック", 2)
    add_bullets(
        doc,
        [
            "人名、数字、日付、製品名が正しい",
            "メールやチャットの宛先が正しい",
            "不要な改行や句読点がない",
            "送信前に全文を読み直した",
        ],
    )

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(18)
    run = closing.add_run("HS Voice 1.2.0  |  社内向け利用マニュアル")
    set_run_font(run, size=9.5, bold=True, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--icon", type=Path, required=True)
    args = parser.parse_args()
    if not args.icon.exists():
        raise FileNotFoundError(args.icon)
    build_manual(args.output, args.icon)


if __name__ == "__main__":
    main()
